#!/usr/bin/env python3
"""Convert the scenario training markdown to a formatted DOCX."""
import re, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

src = r"C:\Users\justi\workspace\prison-policy-ai\officer-training\Scenario-Training-Questions.md"
out = r"C:\Users\justi\workspace\prison-policy-ai\officer-training\New-Officer-Scenario-Training.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('New Officer Scenario Training — 25 Questions', level=0)
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = subtitle.add_run('Based on ADC NCU Policy | NotebookLM-sourced')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

with open(src, 'r', encoding='utf-8') as f:
    lines = f.readlines()

current_section = None
in_answer = False
q_num = 0

for line in lines:
    line = line.rstrip()
    
    # Skip the frontmatter lines
    if line.startswith('#') and line.count('#') == 1 and not line.startswith('##'):
        if '25 Questions' in line or 'Scenario Training' in line:
            continue
    
    # Section headers (## USE OF FORCE, ## PREA, etc.)
    if line.startswith('## ') and not line.startswith('###'):
        heading_text = line.replace('## ', '')
        doc.add_heading(heading_text, level=1)
        current_section = heading_text.strip()
        in_answer = False
        continue
    
    # Q headers (### Q1:)
    if line.startswith('### Q'):
        in_answer = False
        heading = doc.add_heading(line.replace('### ', ''), level=2)
        q_num += 1
        continue
    
    # Scenario header
    if line.startswith('**Scenario:**'):
        in_answer = False
        p = doc.add_paragraph()
        run = p.add_run('Scenario: ')
        run.bold = True
        run.font.size = Pt(11)
        rest = line.replace('**Scenario:**', '').strip()
        if rest:
            p.add_run(rest).font.size = Pt(11)
        continue
    
    # Question header
    if line.startswith('**Question:**'):
        in_answer = False
        p = doc.add_paragraph()
        run = p.add_run('Question: ')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        rest = line.replace('**Question:**', '').strip()
        if rest:
            p.add_run(rest).font.size = Pt(11)
        continue
    
    # Answer header
    if line.startswith('**Answer:**'):
        in_answer = True
        p = doc.add_paragraph()
        run = p.add_run('Answer: ')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 100, 0)
        rest = line.replace('**Answer:**', '').strip()
        if rest.startswith('**') and rest.endswith('**'):
            run2 = p.add_run(rest.strip('*'))
            run2.bold = True
        else:
            p.add_run(rest)
        continue
    
    if in_answer:
        p = doc.add_paragraph()
        if line.startswith('- '):
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(line[2:])
        elif line.startswith('  - '):
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(line[4:])
        elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or \
             line.startswith('4.') or line.startswith('5.') or line.startswith('6.') or \
             line.startswith('7.') or line.startswith('8.') or line.startswith('9.') or \
             line.startswith('10.'):
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(line)
        elif line.strip() == '':
            p.add_run('')
        else:
            # Bold text within answers (policy references etc)
            text = line
            p.add_run(text)
    
    # Empty line between questions
    if line.strip() == '' and not in_answer:
        continue

doc.save(out)
print(f"Saved: {out}")
print(f"Questions: {q_num}")
