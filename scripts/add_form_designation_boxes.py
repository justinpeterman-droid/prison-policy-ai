"""Add the {{box_005}} / {{box_409}} placeholders to the 005 template.

STYLE_RULINGS.md ruling 9: the 005 header carries an Incident Report (005) box
and a Use of Force (409) box, and a use of force ticks BOTH. Until this ran the
two checkbox cells were empty with no placeholder, so filler.py could not tick
either one — no generated 005 had a form designation marked at all.

One-shot: run once, commit the modified .docx, done. Re-running is a no-op.

    PYTHONPATH=. python3 scripts/add_form_designation_boxes.py
"""
from pathlib import Path

from docx import Document

TEMPLATE = Path(__file__).parent.parent / "templates" / "005_template_v3.docx"

# The header table lays each designation out as: [label][checkbox][name], e.g.
# row 0 -> '005' | (bordered empty cell) | 'Incident Report'
#     row, label, placeholder
BOXES = [
    (0, "005", "{{box_005}}"),
    (2, "409", "{{box_409}}"),
]
LABEL_COL = 4
BOX_COL = 5


def add_placeholders() -> int:
    doc = Document(TEMPLATE)
    header = doc.tables[0]
    added = 0

    for row_idx, label, placeholder in BOXES:
        cells = list(header.rows[row_idx].cells)
        # Verify we are writing into the cell we think we are — a template
        # revision that moved these columns must fail loudly, not silently
        # stamp an X into the wrong box.
        actual = cells[LABEL_COL].text.strip()
        if actual != label:
            raise SystemExit(
                f"Expected '{label}' at row {row_idx} col {LABEL_COL}, found "
                f"{actual!r}. The template layout changed — update BOXES.")

        box = cells[BOX_COL]
        if placeholder in box.text:
            print(f"  {placeholder}: already present, skipping")
            continue
        if box.text.strip():
            raise SystemExit(
                f"Checkbox cell for {label} is not empty ({box.text!r}) — "
                f"refusing to overwrite it.")

        para = box.paragraphs[0]
        if para.runs:
            # Reuse the existing run so the cell keeps its Arial sizing and
            # centering rather than inheriting the document default.
            para.runs[0].text = placeholder
        else:
            para.add_run(placeholder)
        print(f"  {placeholder}: added to the {label} box")
        added += 1

    if added:
        doc.save(TEMPLATE)
        print(f"\nSaved {TEMPLATE.name} ({added} placeholder(s) added)")
    else:
        print("\nNothing to do.")
    return added


if __name__ == "__main__":
    add_placeholders()
